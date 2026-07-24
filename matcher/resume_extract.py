"""
resume_extract.py -- pulls raw visible text out of a .docx resume using
python-docx.

Free, local, no API key: python-docx is a pure-Python library that reads
the .docx XML directly, no network call, no signup. This module only
extracts *raw text* -- the `skills` list attached to a ResumeProfile is a
manually curated, honest read of that extracted text by a human (this run),
not automated NER/parsing. Full automated resume field extraction
(spaCy NER, pyresparser) is explicitly out of scope for this increment --
see matcher/README.md "Honest limitations".
"""

from __future__ import annotations

import os
import sys

from docx import Document


def extract_text_from_docx(path: str) -> str:
    """Extracts visible text (paragraphs, then table cells, in document
    order) from a .docx file.

    Raises FileNotFoundError with a clear message if `path` doesn't
    resolve -- per NEXT-BUILD-SPEC.md's builder note, a missing resume
    path is a real blocker to log, never a reason to fabricate resume
    content or silently skip the real-resume dogfood case.
    """
    if not os.path.isfile(path):
        raise FileNotFoundError(
            f"resume .docx not found at {path!r} -- this is a real blocker, "
            "not something to work around by fabricating resume content. "
            "See BUILD-STATUS.md if this was hit during an unattended run."
        )
    doc = Document(path)
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("usage: python resume_extract.py <path-to-resume.docx>", file=sys.stderr)
        raise SystemExit(2)
    print(extract_text_from_docx(sys.argv[1]))
